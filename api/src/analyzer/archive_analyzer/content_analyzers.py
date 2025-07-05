"""Content analyzers for extracting text from various file formats."""

import io
import json
import logging
from typing import List, Optional

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import markdown
except ImportError:
    markdown = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import chardet
except ImportError:
    chardet = None

from .core import FileContent, FileType, ContentAnalyzer

logger = logging.getLogger(__name__)


class TextAnalyzer:
    """Basic text analyzer for plain text files."""

    def can_analyze(self, file_content: FileContent) -> bool:
        """Check if this is a text file."""
        return file_content.file_type in {FileType.TEXT, FileType.CODE}

    def extract_text(self, file_content: FileContent) -> str:
        """Extract text from plain text files."""
        if file_content.content:
            return file_content.content

        # Try to decode with chardet if available
        if chardet and file_content.raw_content:
            detected = chardet.detect(file_content.raw_content)
            encoding = detected.get('encoding', 'utf-8')
            try:
                return file_content.raw_content.decode(encoding)
            except Exception as e:
                logger.warning(f"Failed to decode {file_content.path} with {encoding}: {e}")

        return ""


class PDFAnalyzer:
    """Analyzer for PDF documents."""

    def __init__(self):
        """Initialize PDF analyzer."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF analysis. Install with: pip install PyPDF2")

    def can_analyze(self, file_content: FileContent) -> bool:
        """Check if this is a PDF file."""
        return file_content.path.lower().endswith('.pdf')

    def extract_text(self, file_content: FileContent) -> str:
        """Extract text from PDF files."""
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_content.raw_content))
            text_parts = []

            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num} from {file_content.path}: {e}")

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Failed to analyze PDF {file_content.path}: {e}")
            return ""


class DocxAnalyzer:
    """Analyzer for DOCX documents."""

    def __init__(self):
        """Initialize DOCX analyzer."""
        if Document is None:
            raise ImportError("python-docx is required for DOCX analysis. Install with: pip install python-docx")

    def can_analyze(self, file_content: FileContent) -> bool:
        """Check if this is a DOCX file."""
        return file_content.path.lower().endswith('.docx')

    def extract_text(self, file_content: FileContent) -> str:
        """Extract text from DOCX files."""
        try:
            doc = Document(io.BytesIO(file_content.raw_content))
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Failed to analyze DOCX {file_content.path}: {e}")
            return ""


class MarkdownAnalyzer:
    """Analyzer for Markdown files."""

    def can_analyze(self, file_content: FileContent) -> bool:
        """Check if this is a Markdown file."""
        return file_content.path.lower().endswith(('.md', '.markdown'))

    def extract_text(self, file_content: FileContent) -> str:
        """Extract text from Markdown files."""
        if file_content.content:
            # Remove markdown formatting if the library is available
            if markdown and BeautifulSoup:
                try:
                    html = markdown.markdown(file_content.content)
                    soup = BeautifulSoup(html, 'html.parser')
                    return soup.get_text()
                except Exception as e:
                    logger.warning(f"Failed to parse markdown {file_content.path}: {e}")

            # Fallback: return raw content
            return file_content.content
        return ""


class JSONAnalyzer:
    """Analyzer for JSON files."""

    def can_analyze(self, file_content: FileContent) -> bool:
        """Check if this is a JSON file."""
        return file_content.path.lower().endswith('.json') or file_content.file_type == FileType.DATA

    def extract_text(self, file_content: FileContent) -> str:
        """Extract text from JSON files."""
        if not file_content.content:
            return ""

        try:
            data = json.loads(file_content.content)
            # Extract all string values from the JSON
            return self._extract_strings(data)
        except json.JSONDecodeError:
            # Not valid JSON, return raw content
            return file_content.content

    def _extract_strings(self, obj, strings: Optional[List[str]] = None) -> str:
        """Recursively extract all string values from a JSON object."""
        if strings is None:
            strings = []

        if isinstance(obj, str):
            strings.append(obj)
        elif isinstance(obj, dict):
            for value in obj.values():
                self._extract_strings(value, strings)
        elif isinstance(obj, list):
            for item in obj:
                self._extract_strings(item, strings)

        return ' '.join(strings)


class NotebookAnalyzer:
    """Analyzer for Jupyter Notebook files."""

    def can_analyze(self, file_content: FileContent) -> bool:
        """Check if this is a Jupyter Notebook file."""
        return file_content.path.lower().endswith('.ipynb')

    def extract_text(self, file_content: FileContent) -> str:
        """Extract text from Jupyter Notebook files."""
        if not file_content.content:
            return ""

        try:
            notebook = json.loads(file_content.content)
            text_parts = []

            # Extract code and markdown cells
            for cell in notebook.get('cells', []):
                cell_type = cell.get('cell_type', '')
                source = cell.get('source', [])

                if isinstance(source, list):
                    source = ''.join(source)

                if source:
                    text_parts.append(source)

                # Also extract outputs for code cells
                if cell_type == 'code':
                    for output in cell.get('outputs', []):
                        if 'text' in output:
                            text_parts.append(output['text'])
                        elif 'data' in output and 'text/plain' in output['data']:
                            text_parts.append(output['data']['text/plain'])

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Failed to analyze notebook {file_content.path}: {e}")
            return file_content.content


class ContentAnalyzerFactory:
    """Factory for managing content analyzers."""

    def __init__(self):
        """Initialize the factory with available analyzers."""
        self._analyzers: List[ContentAnalyzer] = []

        # Always available
        self._analyzers.append(TextAnalyzer())
        self._analyzers.append(JSONAnalyzer())
        self._analyzers.append(MarkdownAnalyzer())
        self._analyzers.append(NotebookAnalyzer())

        # Optional analyzers
        try:
            self._analyzers.append(PDFAnalyzer())
        except ImportError:
            logger.warning("PDF support not available. Install PyPDF2.")

        try:
            self._analyzers.append(DocxAnalyzer())
        except ImportError:
            logger.warning("DOCX support not available. Install python-docx.")

    def extract_text(self, file_content: FileContent) -> str:
        """Extract text from file using appropriate analyzer."""
        for analyzer in self._analyzers:
            if analyzer.can_analyze(file_content):
                try:
                    return analyzer.extract_text(file_content)
                except Exception as e:
                    logger.error(f"Analyzer {analyzer.__class__.__name__} failed for {file_content.path}: {e}")

        # Fallback to raw content if available
        return file_content.content or ""